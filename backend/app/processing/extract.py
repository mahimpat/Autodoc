import os, fitz, pytesseract, re
from PIL import Image
from docx import Document as DocxDoc

def _clean_ocr_text(text: str) -> str:
    """Clean up common OCR artifacts and improve text quality"""
    if not text:
        return text
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Fix common OCR mistakes
    # Simple character replacements
    char_replacements = {
        '|': 'I',
        '0': 'O',  # Only if surrounded by letters  
        '5': 'S',  # Only if surrounded by letters
        '1': 'l',  # Only if it makes sense contextually
    }
    
    for old_char, new_char in char_replacements.items():
        text = text.replace(old_char, new_char)
    
    # Regex replacements for cleaning
    regex_replacements = [
        # Remove weird artifacts
        (r'[^\w\s\.,!?()-:;\'\"@#$%&*+=]', ''),
        # Fix spacing around punctuation
        (r'\s+([,.!?;:])', r'\1'),
        (r'([,.!?;:])\s+', r'\1 '),
    ]
    
    for pattern, replacement in regex_replacements:
        text = re.sub(pattern, replacement, text)
    
    # Remove lines that are likely OCR garbage (too short, all caps, weird chars)
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line = line.strip()
        if len(line) < 2:  # Skip very short lines (reduced threshold)
            continue
        if re.match(r'^[^a-zA-Z0-9]*$', line):  # Skip lines with no letters or numbers
            continue
        # Keep short all-caps as they might be important abbreviations or headers
        # if len(line) > 1 and line.isupper() and len(line) < 10:  # Skip short all-caps (likely headers)
        #     continue
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines).strip()

def _extract_pdf_advanced(path: str) -> str:
    """Advanced PDF text extraction with table and layout handling"""
    content_sections = []
    
    try:
        with fitz.open(path) as doc:
            for page_num, page in enumerate(doc):
                page_content = []
                page_content.append(f"\n=== PAGE {page_num + 1} ===")
                
                # Method 1: Try structured text extraction
                try:
                    blocks = page.get_text("dict")
                    structured_text = _parse_pdf_blocks(blocks)
                    if structured_text and len(structured_text.strip()) > 50:
                        page_content.append(structured_text)
                    else:
                        raise Exception("Structured extraction failed")
                except:
                    # Method 2: Fall back to simple text extraction
                    simple_text = page.get_text()
                    if simple_text and simple_text.strip():
                        page_content.append(simple_text)
                
                # Method 3: Extract tables separately
                try:
                    tables = _extract_pdf_tables(page)
                    if tables:
                        page_content.append("\n=== TABLES ===")
                        page_content.extend(tables)
                except:
                    pass
                
                content_sections.append("\n".join(page_content))
        
        return "\n\n".join(content_sections)
    
    except Exception as e:
        # Fallback to basic extraction
        try:
            with fitz.open(path) as doc:
                text = []
                for page in doc:
                    text.append(page.get_text())
                return "\n".join(text)
        except:
            return f"[Could not extract text from PDF: {str(e)}]"

def _parse_pdf_blocks(blocks_dict: dict) -> str:
    """Parse PDF blocks to maintain better structure"""
    content = []
    
    for block in blocks_dict.get("blocks", []):
        if "lines" in block:
            block_text = []
            for line in block["lines"]:
                line_text = []
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        # Preserve some formatting information
                        flags = span.get("flags", 0)
                        if flags & 2**4:  # Bold
                            text = f"**{text}**"
                        if flags & 2**1:  # Italic  
                            text = f"*{text}*"
                        line_text.append(text)
                
                if line_text:
                    block_text.append(" ".join(line_text))
            
            if block_text:
                content.append("\n".join(block_text))
    
    return "\n\n".join(content)

def _extract_pdf_tables(page) -> list:
    """Extract tables from PDF page"""
    tables = []
    
    try:
        # Try to find table-like structures using text positioning
        text_dict = page.get_text("dict")
        
        # Look for aligned text that might be tables
        lines_by_y = {}
        for block in text_dict.get("blocks", []):
            if "lines" in block:
                for line in block["lines"]:
                    y = round(line["bbox"][1])  # y-coordinate
                    if y not in lines_by_y:
                        lines_by_y[y] = []
                    
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        x = span.get("bbox", [0])[0]  # x-coordinate
                        if text:
                            lines_by_y[y].append((x, text))
        
        # Convert to table-like structure
        potential_tables = []
        for y, spans in lines_by_y.items():
            if len(spans) >= 3:  # Potential table row
                spans.sort(key=lambda x: x[0])  # Sort by x-coordinate
                row = " | ".join([span[1] for span in spans])
                potential_tables.append(row)
        
        if potential_tables and len(potential_tables) >= 2:
            tables.append("TABLE:\n" + "\n".join(potential_tables[:10]))  # Limit table size
    
    except Exception:
        pass
    
    return tables

def _extract_docx_advanced(path: str) -> str:
    """Advanced Word document extraction with tables, lists, and formatting"""
    try:
        doc = DocxDoc(path)
        content_sections = []
        
        # Extract document properties/metadata
        if doc.core_properties.title:
            content_sections.append(f"TITLE: {doc.core_properties.title}")
        if doc.core_properties.subject:
            content_sections.append(f"SUBJECT: {doc.core_properties.subject}")
        if doc.core_properties.author:
            content_sections.append(f"AUTHOR: {doc.core_properties.author}")
        
        if content_sections:
            content_sections.append("=" * 50)
        
        # Process document elements in order
        for element in doc.element.body:
            element_text = _process_docx_element(element, doc)
            if element_text and element_text.strip():
                content_sections.append(element_text)
        
        # Extract tables separately (in case they weren't caught above)
        table_content = _extract_docx_tables(doc)
        if table_content:
            content_sections.append("\n=== DOCUMENT TABLES ===")
            content_sections.extend(table_content)
        
        return "\n\n".join(content_sections)
    
    except Exception as e:
        # Fallback to basic extraction
        try:
            doc = DocxDoc(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except:
            return f"[Could not extract text from Word document: {str(e)}]"

def _process_docx_element(element, doc) -> str:
    """Process individual Word document elements"""
    from docx.oxml.ns import qn
    
    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
    
    # Handle paragraphs
    if tag == 'p':
        try:
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if para and para.text.strip():
                text = para.text.strip()
                
                # Check for list items
                if para.style.name.startswith('List'):
                    text = f"• {text}"
                elif para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    text = f"{'#' * int(level if level.isdigit() else 1)} {text}"
                
                return text
        except:
            pass
    
    # Handle tables
    elif tag == 'tbl':
        try:
            for table in doc.tables:
                if table._element == element:
                    return _format_docx_table(table)
        except:
            pass
    
    return ""

def _format_docx_table(table) -> str:
    """Format a Word table as readable text"""
    try:
        table_text = ["=== TABLE ==="]
        
        for i, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_text.append(cell_text)
            
            if any(row_text):  # Skip empty rows
                if i == 0:  # Header row
                    table_text.append(" | ".join(row_text))
                    table_text.append("-" * 50)
                else:
                    table_text.append(" | ".join(row_text))
        
        return "\n".join(table_text)
    except:
        return ""

def _extract_docx_tables(doc) -> list:
    """Extract all tables from Word document"""
    tables = []
    
    try:
        for i, table in enumerate(doc.tables):
            table_content = [f"TABLE {i + 1}:"]
            
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_text.append(cell_text)
                
                if any(row_text):
                    table_content.append(" | ".join(row_text))
            
            if len(table_content) > 1:  # Has content beyond header
                tables.append("\n".join(table_content))
    
    except Exception:
        pass
    
    return tables

SUPPORTED = {'.pdf', '.png', '.jpg', '.jpeg', '.txt', '.md', '.docx', '.doc', '.rtf', '.odt'}

def _cache_path(path: str) -> str:
    base, _ = os.path.splitext(path)
    return base + ".txt"

def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED:
        return f"[UNSUPPORTED FILE TYPE: {ext}]"
    if ext in {'.txt', '.md'}:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    if ext == '.docx':
        return _extract_docx_advanced(path)
    if ext in {'.png', '.jpg', '.jpeg'}:
        img = Image.open(path)
        
        # Try multiple OCR approaches for better handwritten text recognition
        ocr_results = []
        
        # Approach 1: Standard OCR with less restrictive settings
        try:
            config1 = r'--oem 3 --psm 6'
            text1 = pytesseract.image_to_string(img, config=config1)
            if len(text1.strip()) > 5:
                ocr_results.append(("standard", text1.strip()))
        except:
            pass
            
        # Approach 2: Single text block mode (good for handwritten notes)
        try:
            config2 = r'--oem 3 --psm 8'
            text2 = pytesseract.image_to_string(img, config=config2)
            if len(text2.strip()) > 5:
                ocr_results.append(("block", text2.strip()))
        except:
            pass
            
        # Approach 3: Uniform text block (for clean handwriting)
        try:
            config3 = r'--oem 3 --psm 13'
            text3 = pytesseract.image_to_string(img, config=config3)
            if len(text3.strip()) > 5:
                ocr_results.append(("uniform", text3.strip()))
        except:
            pass
        
        # Choose the best result (longest meaningful text)
        if ocr_results:
            best_result = max(ocr_results, key=lambda x: len(x[1]))
            extracted_text = best_result[1]
            
            # Clean up common OCR artifacts
            extracted_text = _clean_ocr_text(extracted_text)
            
            return extracted_text
        
        return "[Could not extract text from image]"
    if ext == '.pdf':
        return _extract_pdf_advanced(path)
    return ""

def extract_text_cached(path: str) -> str:
    cpath = _cache_path(path)
    if os.path.exists(cpath):
        try:
            with open(cpath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            pass
    text = extract_text(path)
    try:
        with open(cpath, 'w', encoding='utf-8') as f:
            f.write(text)
    except Exception:
        pass
    return text

def to_snippets(text: str, min_len: int = 25, max_len: int = 2000, overlap: int = 300):
    """Create comprehensive snippets optimized for source-truth documentation generation.
    
    Since uploaded documents are the source of truth, we need to capture ALL information
    with minimal loss. This function creates overlapping chunks with aggressive coverage.
    """
    if not text or not text.strip():
        return []
    
    text = text.strip()
    
    # Always include the original text for very short content
    if len(text) < min_len:
        return [text] if text else []
    
    # For source-truth documents, we want comprehensive coverage with multiple strategies
    all_chunks = []
    
    # Strategy 1: Preserve document structure - split on headers and sections
    section_chunks = []
    header_pattern = r'\n(===.*?===|\#{1,6}\s+.+|[A-Z][A-Z\s]{10,})\n'
    sections = re.split(header_pattern, text, flags=re.IGNORECASE)
    
    current_section = ""
    for i, section in enumerate(sections):
        if section and section.strip():
            current_section += section + "\n"
            # If section is getting large, create a chunk
            if len(current_section) >= max_len * 0.8:
                if len(current_section.strip()) >= min_len:
                    section_chunks.append(current_section.strip())
                current_section = ""
    
    # Add remaining section content
    if current_section.strip() and len(current_section.strip()) >= min_len:
        section_chunks.append(current_section.strip())
    
    # Strategy 2: Paragraph-based chunking for comprehensive coverage
    para_chunks = []
    paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
    
    # If no clear paragraphs, split on single newlines and group more aggressively
    if len(paragraphs) <= 2:
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if len(lines) > 1:
            paragraphs = []
            current_para = []
            
            for line in lines:
                current_para.append(line)
                para_text = '\n'.join(current_para)
                
                # Create chunks more frequently to capture more information
                should_chunk = (
                    len(para_text) >= min_len * 2 and (
                        line.endswith('.') or line.endswith(':') or 
                        line.endswith('?') or line.endswith('!') or
                        line.endswith(';') or line.endswith(',') or
                        '|' in line or  # Table rows
                        line.isupper() or  # Headers
                        re.match(r'^\d+\.', line.strip()) or  # Numbered lists
                        re.match(r'^[-*•]', line.strip())  # Bullet points
                    )
                )
                
                if should_chunk or len(para_text) >= max_len * 0.7:
                    if len(para_text) >= min_len:
                        paragraphs.append(para_text)
                    current_para = []
            
            # Add remaining content
            if current_para:
                remaining_text = '\n'.join(current_para)
                if len(remaining_text.strip()) >= min_len:
                    paragraphs.append(remaining_text)
    
    # Create overlapping chunks from paragraphs
    current_chunk = ""
    for para in paragraphs:
        test_chunk = (current_chunk + "\n\n" + para).strip() if current_chunk else para
        
        if len(test_chunk) <= max_len:
            current_chunk = test_chunk
        else:
            # Save current chunk if it meets minimum length
            if current_chunk and len(current_chunk) >= min_len:
                para_chunks.append(current_chunk)
            
            # Start new chunk with overlap
            if len(current_chunk) > overlap:
                overlap_text = current_chunk[-overlap:]
                current_chunk = (overlap_text + "\n\n" + para).strip()
            else:
                current_chunk = para
    
    # Add final chunk
    if current_chunk and len(current_chunk) >= min_len:
        para_chunks.append(current_chunk)
    
    # Strategy 3: Sliding window for maximum coverage
    window_chunks = []
    if len(text) > max_len:
        step_size = max_len - overlap
        for i in range(0, len(text), step_size):
            chunk = text[i:i + max_len]
            if len(chunk.strip()) >= min_len:
                window_chunks.append(chunk.strip())
    
    # Strategy 4: Special handling for structured content (tables, lists, code)
    structured_chunks = []
    
    # Extract tables
    table_matches = re.finditer(r'(\|.*\|.*\n){2,}', text, re.MULTILINE)
    for match in table_matches:
        table_text = match.group(0).strip()
        if len(table_text) >= min_len:
            structured_chunks.append("TABLE:\n" + table_text)
    
    # Extract code blocks
    code_matches = re.finditer(r'```[\s\S]*?```|`[^`\n]+`', text)
    for match in code_matches:
        code_text = match.group(0).strip()
        if len(code_text) >= min_len:
            structured_chunks.append("CODE:\n" + code_text)
    
    # Extract numbered/bulleted lists
    list_matches = re.finditer(r'(^[ \t]*[\d\w\-*•][\.\)]\s+.*\n){2,}', text, re.MULTILINE)
    for match in list_matches:
        list_text = match.group(0).strip()
        if len(list_text) >= min_len:
            structured_chunks.append("LIST:\n" + list_text)
    
    # Combine all strategies and deduplicate
    all_chunks = section_chunks + para_chunks + window_chunks + structured_chunks
    
    # Remove duplicates while preserving order
    seen = set()
    unique_chunks = []
    for chunk in all_chunks:
        # Use a hash of the first 100 chars to detect near-duplicates
        chunk_key = chunk[:100].strip()
        if chunk_key not in seen and len(chunk) >= min_len:
            unique_chunks.append(chunk)
            seen.add(chunk_key)
    
    # If no good chunks were created, fall back to simple chunking
    if not unique_chunks:
        # Very aggressive fallback - chunk by sentences with minimal requirements
        sentences = re.split(r'[.!?]+\s+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if sentences:
            current_chunk = ""
            for sentence in sentences:
                test_chunk = (current_chunk + ". " + sentence).strip() if current_chunk else sentence
                
                if len(test_chunk) <= max_len:
                    current_chunk = test_chunk
                else:
                    if current_chunk:
                        unique_chunks.append(current_chunk + ".")
                    current_chunk = sentence
            
            if current_chunk:
                unique_chunks.append(current_chunk)
        
        # Ultimate fallback - just use the original text in chunks
        if not unique_chunks:
            for i in range(0, len(text), max_len - overlap):
                chunk = text[i:i + max_len]
                if chunk.strip():
                    unique_chunks.append(chunk.strip())
    
    # Ensure we have at least something if the original text exists
    if not unique_chunks and text.strip():
        unique_chunks = [text]
    
    print(f"DEBUG CHUNKING: Created {len(unique_chunks)} chunks from {len(text)} chars using comprehensive strategy")
    return unique_chunks
